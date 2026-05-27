# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn

doc = Document(r'F:\26毕设单\基于Java的志愿服务活动管理系统的设计与实现\基于Java的志愿服务活动管理系统的设计与实现-初稿.docx')

with open(r'F:\26毕设2\协作编辑系统\inspect_result.txt', 'w', encoding='utf-8') as f:
    # 1. 摘要段落字数
    f.write("=== 摘要内容 ===\n")
    for i in [19, 20, 21]:
        text = doc.paragraphs[i].text
        f.write(f"段落[{i}]: 字数={len(text)}\n{text}\n\n")
    
    # 2. 2.5 JWT段落详细 - 检查runs中的字间距
    f.write("\n=== 2.5 JWT段落(72) runs详情 ===\n")
    p72 = doc.paragraphs[72]
    for j, run in enumerate(p72.runs):
        spacing = None
        rpr = run._element.find(qn('w:rPr'))
        if rpr is not None:
            sp = rpr.find(qn('w:spacing'))
            if sp is not None:
                spacing = sp.get(qn('w:val'))
        f.write(f"  run[{j}]: spacing={spacing} text='{run.text}'\n")
    
    # 3. Heading 1 样式详情
    f.write("\n=== Heading 1 段落样式 ===\n")
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name == 'Heading 1':
            f.write(f"  段落[{i}]: '{p.text}'\n")
            for j, run in enumerate(p.runs):
                font = run.font
                f.write(f"    run[{j}]: bold={font.bold}, color={font.color.rgb if font.color and font.color.rgb else 'None'}, size={font.size}, text='{run.text}'\n")
    
    # 4. Heading 2 样式详情
    f.write("\n=== Heading 2 段落样式(前5个) ===\n")
    count = 0
    for i, p in enumerate(doc.paragraphs):
        if p.style and p.style.name == 'Heading 2':
            f.write(f"  段落[{i}]: '{p.text}'\n")
            for j, run in enumerate(p.runs):
                font = run.font
                f.write(f"    run[{j}]: bold={font.bold}, color={font.color.rgb if font.color and font.color.rgb else 'None'}, size={font.size}, text='{run.text}'\n")
            count += 1
            if count >= 5:
                break
    
    # 5. 表格信息
    f.write(f"\n=== 表格总数: {len(doc.tables)} ===\n")
    for t_idx, table in enumerate(doc.tables):
        f.write(f"\n表格[{t_idx}]: 行数={len(table.rows)}, 列数={len(table.columns)}\n")
        # 检查表格边框
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            borders = tblPr.find(qn('w:tblBorders'))
            if borders is not None:
                for child in borders:
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    val = child.get(qn('w:val'))
                    sz = child.get(qn('w:sz'))
                    f.write(f"  边框 {tag}: val={val}, sz={sz}\n")
        # 打印前两行内容
        for r_idx, row in enumerate(table.rows):
            if r_idx < 2:
                cells = [cell.text[:20] for cell in row.cells]
                f.write(f"  行[{r_idx}]: {cells}\n")
    
    # 6. 致谢内容
    f.write("\n=== 致谢内容 ===\n")
    for i in range(309, 320):
        if i < len(doc.paragraphs):
            text = doc.paragraphs[i].text
            if text.strip():
                f.write(f"段落[{i}]: {text}\n\n")
    
    # 7. 参考文献
    f.write("\n=== 参考文献 ===\n")
    for i in range(289, 310):
        if i < len(doc.paragraphs):
            text = doc.paragraphs[i].text
            if text.strip():
                f.write(f"段落[{i}]: {text}\n")
    
    # 8. 3.1 功能需求 段落78-79
    f.write("\n=== 3.1 功能需求分析(78-79) ===\n")
    for i in [77, 78, 79]:
        text = doc.paragraphs[i].text
        f.write(f"段落[{i}]: {text}\n\n")
