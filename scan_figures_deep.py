from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os

doc = Document(r'F:\26毕设单\单片机智慧养老家居\修后论文_8章新版_修改版.docx')

output = []

output.append('=== 深度扫描：每个图注前后5段的上下文 ===\n')

all_paragraphs = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs)]

fig_indices = []
for i, text in all_paragraphs:
    if text and (text.startswith('图') or text.startswith('Figure')):
        fig_indices.append(i)

for fig_idx in fig_indices:
    fig_text = all_paragraphs[fig_idx][1] if fig_idx < len(all_paragraphs) else ''
    output.append(f'--- {fig_text[:100]} ---')
    
    start = max(0, fig_idx - 3)
    end = min(len(all_paragraphs), fig_idx + 3)
    
    for j in range(start, end):
        idx, text = all_paragraphs[j]
        if text:
            marker = ' >>>' if j == fig_idx else '    '
            output.append(f'{marker} [段落{idx}] {text[:200]}')
    output.append('')

output.append('\n=== 检查是否有图片的alt文本或描述含不相关硬件 ===')
bad_kw = ['毫米波', 'HLK-LD2410', 'HC-SR501', '红外', '人体感应', 'PIR', 'radar', '雷达', 'LD2410', 'SR501', '热释电']

from docx.oxml.ns import qn
img_count = 0
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        drawing_elements = run._element.findall(qn('w:drawing'))
        for drawing in drawing_elements:
            img_count += 1
            desc_elements = drawing.findall('.//' + qn('wp:docPr'))
            for desc in desc_elements:
                name = desc.get('name', '')
                descr = desc.get('descr', '')
                if name or descr:
                    has_bad = any(kw in name or kw in descr for kw in bad_kw)
                    flag = ' WARNING' if has_bad else ''
                    output.append(f'[段落{i}] 图片name="{name}" descr="{descr[:100]}"{flag}')

output.append(f'\n论文中共有 {img_count} 张嵌入图片')

output.append('\n=== 特别关注：硬件相关图（图5-x系列）的完整上下文 ===')
for fig_idx in fig_indices:
    fig_text = all_paragraphs[fig_idx][1] if fig_idx < len(all_paragraphs) else ''
    if '5-' in fig_text or '硬件' in fig_text:
        output.append(f'\n=== {fig_text[:100]} 前后段落 ===')
        start = max(0, fig_idx - 5)
        end = min(len(all_paragraphs), fig_idx + 5)
        for j in range(start, end):
            idx, text = all_paragraphs[j]
            if text:
                marker = '>>>' if j == fig_idx else '   '
                output.append(f'{marker} [{idx}] {text[:300]}')

result = '\n'.join(output)
with open(r'f:\26毕设2\协作编辑系统\scan_deep_result.txt', 'w', encoding='utf-8') as f:
    f.write(result)

print('Done. Deep scan results written.')
