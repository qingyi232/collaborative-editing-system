import sys
from docx import Document

doc = Document(r'F:\26毕设单\单片机智慧养老家居\修后论文_8章新版_修改版.docx')

bad_keywords = ['毫米波', 'HLK-LD2410', 'HC-SR501', '红外', '人体感应', 'PIR', 'radar', '雷达', 'LD2410', 'SR501', '热释电']

output = []

output.append('=== 所有含"图"字且含不相关硬件的段落 ===')
fig_count = 0
fig_bad_count = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    if '图' in text:
        fig_count += 1
        for kw in bad_keywords:
            if kw in text:
                fig_bad_count += 1
                output.append(f'[段落 {i}] WARNING "{kw}": {text[:250]}')
                break

output.append(f'\n含"图"字段落总数: {fig_count}, 其中含不相关硬件: {fig_bad_count}')

output.append('\n=== 所有仍含不相关硬件关键词的段落 ===')
found_bad = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    for kw in bad_keywords:
        if kw in text:
            found_bad.append((i, kw, text[:250]))
            break

for idx, kw, txt in found_bad:
    output.append(f'[段落 {idx}] 关键词="{kw}": {txt}')

output.append(f'\n总计 {len(found_bad)} 个段落仍含不相关硬件关键词')

output.append('\n=== 检查表格中是否残留不相关硬件 ===')
table_bad = 0
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                for kw in bad_keywords:
                    if kw in text:
                        table_bad += 1
                        output.append(f'[表格{ti} 行{ri} 列{ci}] WARNING "{kw}": {text[:200]}')
                        break

output.append(f'\n表格中含不相关硬件: {table_bad} 处')

output.append('\n=== 列出所有图注（以"图"开头的段落） ===')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and (text.startswith('图') or text.startswith('Figure')):
        output.append(f'[段落 {i}] 图注: {text[:200]}')

result = '\n'.join(output)
with open(r'f:\26毕设2\协作编辑系统\scan_result.txt', 'w', encoding='utf-8') as f:
    f.write(result)

print('Done. Results written to scan_result.txt')
