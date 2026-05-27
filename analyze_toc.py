# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from lxml import etree

doc = Document(r'F:\26毕设单\基于Java的文学兴趣社区平台的设计与实现\2201010573-王子涵-基于Java的文学兴趣社区平台的设计与实现(1).docx')

nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

print('=== ALL SECTIONS FOOTER DETAILS ===')
for sec_idx in range(len(doc.sections)):
    section = doc.sections[sec_idx]
    footer = section.footer
    print(f'\nSection {sec_idx}: linked={footer.is_linked_to_previous}')
    for i, p in enumerate(footer.paragraphs):
        print(f'  Para {i}: text="{p.text}" align={p.paragraph_format.alignment}')
        xml = etree.tostring(p._element, pretty_print=True).decode()
        if 'fldChar' in xml or 'instrText' in xml or 'PAGE' in xml:
            print(f'    [Has page number field]')
        if not p.text.strip() and 'fldChar' not in xml:
            print(f'    [EMPTY - no page number]')

print('\n=== TOC AREA PARAGRAPHS ===')
for i in range(38, 45):
    p = doc.paragraphs[i]
    print(f'[{i}] style={p.style.name} | "{p.text[:60]}"')

print('\n=== SECTION 3 FOOTER XML ===')
sec3 = doc.sections[3]
footer3_xml = etree.tostring(sec3.footer._element, pretty_print=True).decode()
print(footer3_xml[:2000])
